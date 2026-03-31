"""Tests for office document conversions (skipped if deps not installed)."""

from __future__ import annotations

import importlib
import json
from pathlib import Path

import pytest

from cvt.backends.base import _try_import
from cvt.converter import convert_file


def _has(*pkgs: str) -> bool:
    return all(_try_import(p) for p in pkgs)


# ---------------------------------------------------------------------------
# XLSX conversions
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_xlsx(tmp_path):
    if not _has("openpyxl"):
        pytest.skip("openpyxl not installed")
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["id", "name", "score"])
    ws.append([1, "Alice", 92.5])
    ws.append([2, "Bob", 78.0])
    path = tmp_path / "sample.xlsx"
    wb.save(str(path))
    return path


@pytest.mark.skipif(not _has("openpyxl"), reason="openpyxl not installed")
def test_xlsx_to_csv(sample_xlsx, tmp_path):
    out = tmp_path / "out.csv"
    convert_file(sample_xlsx, out)
    content = out.read_text()
    assert "Alice" in content
    assert "id" in content


@pytest.mark.skipif(not _has("openpyxl"), reason="openpyxl not installed")
def test_xlsx_to_json(sample_xlsx, tmp_path):
    out = tmp_path / "out.json"
    convert_file(sample_xlsx, out)
    data = json.loads(out.read_text())
    assert isinstance(data, list)
    assert data[0]["name"] == "Alice"


@pytest.mark.skipif(not _has("openpyxl"), reason="openpyxl not installed")
def test_csv_to_xlsx(tmp_path):
    csv_file = tmp_path / "sample.csv"
    csv_file.write_text("id,name\n1,Alice\n2,Bob\n", encoding="utf-8")
    out = tmp_path / "out.xlsx"
    convert_file(csv_file, out)
    import openpyxl
    wb = openpyxl.load_workbook(str(out))
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    assert rows[0] == ("id", "name")
    assert rows[1][1] == "Alice"


# ---------------------------------------------------------------------------
# DOCX conversions
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_docx(tmp_path):
    if not _has("docx"):
        pytest.skip("python-docx not installed")
    from docx import Document
    doc = Document()
    doc.add_heading("CVT Test Document", 0)
    doc.add_paragraph("This is a test paragraph for conversion.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.mark.skipif(not _has("mammoth"), reason="mammoth not installed")
def test_docx_to_html(sample_docx, tmp_path):
    out = tmp_path / "out.html"
    convert_file(sample_docx, out)
    content = out.read_text()
    assert "<" in content


@pytest.mark.skipif(not _has("mammoth"), reason="mammoth not installed")
def test_docx_to_markdown(sample_docx, tmp_path):
    out = tmp_path / "out.md"
    convert_file(sample_docx, out)
    content = out.read_text()
    assert len(content) > 0


@pytest.mark.skipif(not _has("mammoth"), reason="mammoth not installed")
def test_docx_to_txt(sample_docx, tmp_path):
    out = tmp_path / "out.txt"
    convert_file(sample_docx, out)
    content = out.read_text()
    assert "test" in content.lower()


# ---------------------------------------------------------------------------
# PDF → TXT
# ---------------------------------------------------------------------------

@pytest.mark.skipif(not _has("pdfminer"), reason="pdfminer.six not installed")
def test_pdf_to_txt(tmp_path):
    # Create a minimal PDF using weasyprint if available, else skip
    if not _has("weasyprint"):
        pytest.skip("weasyprint not installed (needed to create test PDF)")
    from weasyprint import HTML
    pdf_path = tmp_path / "sample.pdf"
    HTML(string="<html><body><p>Hello from CVT</p></body></html>").write_pdf(str(pdf_path))
    out = tmp_path / "out.txt"
    convert_file(pdf_path, out)
    content = out.read_text()
    assert "Hello" in content or len(content) > 0
